# -*- coding: utf-8 -*-
"""
Dựng file báo cáo Word hoàn chỉnh từ mẫu của giảng viên.

Giữ nguyên trang bìa, cấu trúc phân đoạn, footer đánh số trang (i, ii, iii cho
phần đầu — 1, 2, 3 cho phần thân) và trường mục lục tự động của mẫu; chỉ chuẩn
hóa định dạng và chèn nội dung vào từng mục.
"""
import os
import sys
import copy

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import noi_dung as ND

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
HINH = os.path.join(BASE, "hinh-anh")
MAU = os.path.join(BASE, "nguon", "mau_goc.docx")
OUT = os.path.join(BASE, "BÁO CÁO - Thương mại điện tử - Lumina Tutors.docx")

FONT = "Times New Roman"
CO_CHU = Pt(13)          # cỡ chữ thân bài
THUT_DAU_DONG = Cm(1.0)  # thụt lề dòng đầu đoạn văn

doc = Document(MAU)
body = doc.element.body


# ═══════════════════════════════════════════════════════════════════════════════
#  TIỆN ÍCH XML
# ═══════════════════════════════════════════════════════════════════════════════

def el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k.replace("_", ":")), v)
    return e


def dat_font(run, size=CO_CHU, bold=None, italic=None, color=None):
    run.font.name = FONT
    run.font.size = size
    rpr = run._r.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), FONT)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def them_field(paragraph, instr, cached="", size=CO_CHU, italic=False):
    """Chèn một trường (field) Word vào cuối đoạn văn."""
    def _run(child):
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rf = OxmlElement("w:rFonts")
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(a), FONT)
        rpr.append(rf)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size.pt * 2))); rpr.append(sz)
        szcs = OxmlElement("w:szCs"); szcs.set(qn("w:val"), str(int(size.pt * 2))); rpr.append(szcs)
        if italic:
            rpr.append(OxmlElement("w:i"))
        r.append(rpr)
        r.append(child)
        paragraph._p.append(r)

    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin"); _run(b)
    t = OxmlElement("w:instrText"); t.set(qn("xml:space"), "preserve"); t.text = instr; _run(t)
    s = OxmlElement("w:fldChar"); s.set(qn("w:fldCharType"), "separate"); _run(s)
    c = OxmlElement("w:t"); c.set(qn("xml:space"), "preserve"); c.text = cached; _run(c)
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end"); _run(e)


def to_mau_o(cell, hex_color):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def vien_o(cell, sz=6, color="808080", styles=("top", "left", "bottom", "right"), val="single"):
    tcpr = cell._tc.get_or_add_tcPr()
    borders = tcpr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcpr.append(borders)
    for s in styles:
        b = OxmlElement(f"w:{s}")
        b.set(qn("w:val"), val)
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:color"), color)
        borders.append(b)


def lap_lai_hang_tieu_de(row):
    trpr = row._tr.get_or_add_trPr()
    trpr.append(OxmlElement("w:tblHeader"))


def khong_tach_hang(row):
    trpr = row._tr.get_or_add_trPr()
    trpr.append(OxmlElement("w:cantSplit"))


# ═══════════════════════════════════════════════════════════════════════════════
#  CHUẨN HÓA ĐỊNH DẠNG TOÀN CỤC
# ═══════════════════════════════════════════════════════════════════════════════

def chuan_hoa_styles():
    st = doc.styles

    # -- Normal: Times New Roman 13pt, canh đều, giãn dòng 1.5 --------------
    n = st["Normal"]
    n.font.name = FONT
    n.font.size = CO_CHU
    rpr = n.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), FONT)
    pf = n.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)

    # -- Heading 1/2/3: đồng bộ font, cỡ, khoảng cách ------------------------
    for name, size, bold, italic, align in (
        ("Heading 1", Pt(14), True, False, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", Pt(13), True, False, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", Pt(13), True, True, WD_ALIGN_PARAGRAPH.LEFT),
    ):
        s = st[name]
        s.font.name = FONT
        s.font.size = size
        s.font.bold = bold
        s.font.italic = italic
        s.font.color.rgb = RGBColor(0, 0, 0)
        r = s.element.get_or_add_rPr()
        f = r.find(qn("w:rFonts"))
        if f is None:
            f = OxmlElement("w:rFonts"); r.insert(0, f)
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            f.set(qn(a), FONT)
        s.paragraph_format.alignment = align
        s.paragraph_format.line_spacing = 1.5
        s.paragraph_format.space_before = Pt(12)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.keep_with_next = True

    # -- Caption: TNR 12pt nghiêng, canh giữa, màu đen ----------------------
    c = st["Caption"]
    c.font.name = FONT
    c.font.size = Pt(12)
    c.font.italic = True
    c.font.bold = False
    c.font.color.rgb = RGBColor(0, 0, 0)
    r = c.element.get_or_add_rPr()
    f = r.find(qn("w:rFonts"))
    if f is None:
        f = OxmlElement("w:rFonts"); r.insert(0, f)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        f.set(qn(a), FONT)
    c.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.line_spacing = 1.0
    c.paragraph_format.space_before = Pt(6)
    c.paragraph_format.space_after = Pt(6)

    # -- ListParagraph: bỏ canh giữa mặc định --------------------------------
    lp = st["List Paragraph"]
    lp.paragraph_format.line_spacing = 1.5
    lp.paragraph_format.space_after = Pt(3)


def chuan_hoa_trang():
    """Khổ A4, giữ nguyên lề của mẫu (trên/dưới 2.5cm, trái 3cm, phải 2cm)."""
    for s in doc.sections:
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3.0)
        s.right_margin = Cm(2.0)


def ep_font_toan_tai_lieu():
    """Thay mọi tham chiếu font theo theme (Aptos/Calibri) bằng Times New Roman
    tường minh, ở thân bài, chân trang, đầu trang và style mặc định."""
    THEME_ATTRS = ("w:asciiTheme", "w:hAnsiTheme", "w:cstheme", "w:eastAsiaTheme")
    phan = [doc.element.body, doc.styles.element]
    for s in doc.sections:
        for p in (s.footer, s.header, s.even_page_footer, s.even_page_header,
                  s.first_page_footer, s.first_page_header):
            try:
                phan.append(p._element)
            except Exception:
                pass

    for goc in phan:
        for rf in goc.iter(qn("w:rFonts")):
            for a in THEME_ATTRS:
                if rf.get(qn(a)) is not None:
                    del rf.attrib[qn(a)]
            for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rf.set(qn(a), FONT)

    # Font mặc định của tài liệu
    rpr_def = doc.styles.element.find(qn("w:docDefaults"))
    if rpr_def is not None:
        rpr = rpr_def.find(qn("w:rPrDefault"))
        if rpr is not None:
            r = rpr.find(qn("w:rPr"))
            if r is None:
                r = OxmlElement("w:rPr"); rpr.append(r)
            rf = r.find(qn("w:rFonts"))
            if rf is None:
                rf = OxmlElement("w:rFonts"); r.insert(0, rf)
            for a in THEME_ATTRS:
                if rf.get(qn(a)) is not None:
                    del rf.attrib[qn(a)]
            for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
                rf.set(qn(a), FONT)

    # Chân trang: cỡ chữ 13, canh phải giữ nguyên theo mẫu
    for s in doc.sections:
        for ft in (s.footer, s.even_page_footer, s.first_page_footer):
            try:
                for p in ft.paragraphs:
                    for r in p.runs:
                        r.font.size = CO_CHU
            except Exception:
                pass


def bat_cap_nhat_field():
    """Yêu cầu Word cập nhật toàn bộ trường (mục lục, danh mục) khi mở file."""
    settings = doc.settings.element
    for t in settings.findall(qn("w:updateFields")):
        settings.remove(t)
    uf = OxmlElement("w:updateFields")
    uf.set(qn("w:val"), "true")
    settings.append(uf)


# ═══════════════════════════════════════════════════════════════════════════════
#  DỌN DẸP MẪU
# ═══════════════════════════════════════════════════════════════════════════════

def co_sectpr(p_el):
    ppr = p_el.find(qn("w:pPr"))
    return ppr is not None and ppr.find(qn("w:sectPr")) is not None


def lay_style(p_el):
    ppr = p_el.find(qn("w:pPr"))
    if ppr is None:
        return "Normal"
    st = ppr.find(qn("w:pStyle"))
    return st.get(qn("w:val")) if st is not None else "Normal"


def lay_text(p_el):
    return "".join(t.text or "" for t in p_el.iter(qn("w:t")))


def don_dep():
    """Giữ lại khung xương của mẫu (tiêu đề, dấu ngắt phân đoạn, trường mục lục,
    ba bảng có sẵn) và xóa toàn bộ nội dung đề cương mẫu để thay bằng nội dung
    thật."""
    con = list(body)

    # Xác định khoảng chứa trường mục lục: từ tiêu đề "MỤC LỤC" tới tiêu đề
    # "DANH MỤC CÁC KÝ HIỆU"
    i_mucluc = i_danhmuc = None
    for i, ch in enumerate(con):
        if ch.tag != qn("w:p") or not lay_style(ch).startswith("Heading"):
            continue
        t = lay_text(ch).strip()
        if t == "MỤC LỤC" and i_mucluc is None:
            i_mucluc = i
        if t.startswith("DANH MỤC CÁC KÝ HIỆU") and i_danhmuc is None:
            i_danhmuc = i
    assert i_mucluc is not None and i_danhmuc is not None, "Không xác định được vùng mục lục"

    # Đoạn nào nằm trong một trường Word đang mở thì tuyệt đối không được xóa,
    # nếu không cấu trúc begin/separate/end của trường sẽ hỏng.
    trong_field = set()
    do_sau = 0
    for i, ch in enumerate(con):
        if ch.tag != qn("w:p"):
            continue
        co_field = False
        for fc in ch.iter(qn("w:fldChar")):
            co_field = True
            loai = fc.get(qn("w:fldCharType"))
            if loai == "begin":
                do_sau += 1
            elif loai == "end":
                do_sau = max(0, do_sau - 1)
        if co_field or do_sau > 0:
            trong_field.add(i)

    for i, ch in enumerate(con):
        if ch.tag != qn("w:p"):
            continue                      # giữ nguyên các bảng của mẫu
        if lay_style(ch).startswith("Heading"):
            continue                      # giữ tiêu đề
        if co_sectpr(ch):
            continue                      # giữ dấu ngắt phân đoạn
        if i in trong_field:
            continue                      # giữ nguyên trường mục lục tự động
        if i_mucluc < i < i_danhmuc:
            continue
        ch.getparent().remove(ch)


def sua_giao_dien_muc_luc():
    """Chuẩn hóa style mục lục: bỏ chữ in hoa toàn phần, tăng cỡ chữ, đặt lại
    vị trí tab dấu chấm cho khớp lề trang A4."""
    tab_pos = str(int((21.0 - 3.0 - 2.0) * 567))   # bề rộng vùng chữ, đơn vị twip
    for lvl, size, bold, italic, indent in (
        ("toc 1", Pt(13), True, False, 0),
        ("toc 2", Pt(13), False, False, 340),
        ("toc 3", Pt(13), False, True, 680),
    ):
        s = doc.styles[lvl]
        el_s = s.element
        rpr = el_s.get_or_add_rPr()
        for tag in ("w:caps", "w:smallCaps", "w:b", "w:bCs", "w:i", "w:iCs", "w:sz", "w:szCs"):
            for e in rpr.findall(qn(tag)):
                rpr.remove(e)
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts"); rpr.insert(0, rf)
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(a), FONT)
        s.font.size = size
        s.font.bold = bold
        s.font.italic = italic

        ppr = el_s.get_or_add_pPr()
        for e in ppr.findall(qn("w:tabs")):
            ppr.remove(e)
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:leader"), "dot")
        tab.set(qn("w:pos"), tab_pos)
        tabs.append(tab)
        ppr.insert(0, tabs)
        s.paragraph_format.left_indent = Emu(int(indent * 635))
        s.paragraph_format.line_spacing = 1.3
        s.paragraph_format.space_before = Pt(3)
        s.paragraph_format.space_after = Pt(3)


def sua_giao_dien_trang_bia():
    """Trang bìa được thiết kế cho giãn dòng đơn — khôi phục lại sau khi style
    Normal đổi sang giãn dòng 1.5."""
    cell = doc.tables[0].cell(0, 0)
    for p in cell.paragraphs:
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Ngắt trang giữa bìa ngoài và bìa lót
    lan = 0
    for p in cell.paragraphs:
        if p.text.strip().startswith("BỘ GIÁO DỤC VÀ ĐÀO TẠO"):
            lan += 1
            if lan == 2:
                r = OxmlElement("w:r")
                br = OxmlElement("w:br")
                br.set(qn("w:type"), "page")
                r.append(br)
                p._p.insert(0, r)
                ppr = p._p.find(qn("w:pPr"))
                if ppr is not None:
                    p._p.remove(ppr)
                    p._p.insert(0, ppr)
                break


# ═══════════════════════════════════════════════════════════════════════════════
#  CON TRỎ CHÈN NỘI DUNG
# ═══════════════════════════════════════════════════════════════════════════════

class ConTro:
    """Chèn phần tử mới ngay sau phần tử hiện tại rồi dời con trỏ tới đó."""

    def __init__(self, element):
        self.el = element

    def them(self, element):
        self.el.addnext(element)
        self.el = element
        return element


# Bộ đếm số thứ tự hình/bảng theo chương (chỉ để ghi nhãn cố định trong text)
class DemSo:
    def __init__(self):
        self.chuong = 1
        self.hinh_dau_chuong = True
        self.bang_dau_chuong = True

    def sang_chuong(self, n):
        self.chuong = n
        self.hinh_dau_chuong = True
        self.bang_dau_chuong = True


dem = DemSo()


def p_moi(cursor, style=None):
    p = doc.add_paragraph(style=style)
    cursor.them(p._p)
    return p


def them_doan(cursor, text, thut=True):
    p = p_moi(cursor)
    r = p.add_run(text)
    dat_font(r)
    p.paragraph_format.first_line_indent = THUT_DAU_DONG if thut else None
    return p


def them_gach_dau_dong(cursor, items):
    for it in items:
        p = p_moi(cursor, style="List Paragraph")
        r = p.add_run(it)
        dat_font(r)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        ppr = p._p.get_or_add_pPr()
        numpr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0"); numpr.append(ilvl)
        nid = OxmlElement("w:numId"); nid.set(qn("w:val"), "2"); numpr.append(nid)
        ppr.insert(0, numpr)


def them_chu_thich(cursor, loai, text, tren_bang=False):
    """loai = 'Hình' hoặc 'Bảng'. Dùng trường SEQ để Word tự đánh số và gom vào
    danh mục hình/bảng."""
    seq_id = "Hinh" if loai == "Hình" else "Bang"
    dau_chuong = dem.hinh_dau_chuong if loai == "Hình" else dem.bang_dau_chuong

    p = p_moi(cursor, style="Caption")
    r = p.add_run(f"{loai} {dem.chuong}.")
    dat_font(r, size=Pt(12), italic=True, bold=False)
    instr = f" SEQ {seq_id} \\* ARABIC "
    if dau_chuong:
        instr = f" SEQ {seq_id} \\* ARABIC \\r 1 "
    them_field(p, instr, cached="1", size=Pt(12), italic=True)
    r2 = p.add_run(f". {text}")
    dat_font(r2, size=Pt(12), italic=True, bold=False)
    p.paragraph_format.keep_with_next = tren_bang

    if loai == "Hình":
        dem.hinh_dau_chuong = False
    else:
        dem.bang_dau_chuong = False
    return p


def them_hinh(cursor, ten_file, chu_thich, rong_cm):
    duong_dan = os.path.join(HINH, ten_file)
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(duong_dan, width=Cm(rong_cm))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    cursor.them(p._p)
    them_chu_thich(cursor, "Hình", chu_thich)


def them_khung_cho_anh(cursor, chu_thich, cao_cm):
    """Khung chừa chỗ dán ảnh chụp màn hình."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    cell = t.cell(0, 0)
    cell.width = Cm(15.0)
    t.columns[0].width = Cm(15.0)
    vien_o(cell, sz=6, color="9E9E9E", val="dashed")
    to_mau_o(cell, "FAFAFA")
    cell.vertical_alignment = 1  # WD_ALIGN_VERTICAL.CENTER
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[ Chèn ảnh chụp màn hình: {chu_thich} ]")
    dat_font(r, size=Pt(12), italic=True, color=RGBColor(0x80, 0x80, 0x80))
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    # chiều cao tối thiểu của hàng
    trpr = t.rows[0]._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(cao_cm * 567)))
    h.set(qn("w:hRule"), "atLeast")
    trpr.append(h)
    khong_tach_hang(t.rows[0])
    cursor.them(t._tbl)
    them_chu_thich(cursor, "Hình", chu_thich)
    # đoạn trống sau bảng để tách khỏi nội dung kế tiếp
    return t


def them_bang(cursor, chu_thich, headers, rows, widths_cm):
    them_chu_thich(cursor, "Bảng", chu_thich, tren_bang=True)

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    tblpr = t._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblpr.append(layout)

    def dat_do_rong(row):
        for i, w in enumerate(widths_cm):
            row.cells[i].width = Cm(w)

    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        dat_font(r, size=Pt(12), bold=True)
        to_mau_o(c, "EDF0F5")
    dat_do_rong(hdr)
    lap_lai_hang_tieu_de(hdr)
    khong_tach_hang(hdr)

    for data in rows:
        row = t.add_row()
        for i, val in enumerate(data):
            c = row.cells[i]
            c.text = ""
            dong = str(val).split("\n")
            for j, d in enumerate(dong):
                p = c.paragraphs[0] if j == 0 else c.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(d)
                dat_font(r, size=Pt(12))
        dat_do_rong(row)
        khong_tach_hang(row)
    for i, w in enumerate(widths_cm):
        t.columns[i].width = Cm(w)

    cursor.them(t._tbl)

    # đoạn trống sau bảng
    p = p_moi(cursor)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("")
    dat_font(r, size=Pt(6))
    return t


def do_khoi(cursor, khoi):
    for b in khoi:
        loai = b[0]
        if loai == "p":
            them_doan(cursor, b[1])
        elif loai == "b":
            them_gach_dau_dong(cursor, b[1])
        elif loai == "fig":
            them_hinh(cursor, b[1], b[2], b[3])
        elif loai == "ph":
            them_khung_cho_anh(cursor, b[1], b[2])
        elif loai == "tbl":
            them_bang(cursor, b[1], b[2], b[3], b[4])
        else:
            raise ValueError(f"Khối không hợp lệ: {loai}")


# ═══════════════════════════════════════════════════════════════════════════════
#  TÌM TIÊU ĐỀ
# ═══════════════════════════════════════════════════════════════════════════════

def tim_heading(prefix):
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and p.text.strip().startswith(prefix):
            return p
    raise KeyError(f"Không tìm thấy tiêu đề bắt đầu bằng: {prefix!r}")


def chen_sau(prefix, khoi, chuong=None):
    if chuong is not None:
        dem.sang_chuong(chuong)
    h = tim_heading(prefix)
    do_khoi(ConTro(h._p), khoi)


# ═══════════════════════════════════════════════════════════════════════════════
#  TRANG BÌA
# ═══════════════════════════════════════════════════════════════════════════════

def sua_trang_bia():
    """Điền tên đề tài vào cả bìa ngoài lẫn bìa lót."""
    bia = doc.tables[0]
    for p in bia.cell(0, 0).paragraphs:
        if "Tên đề tài" not in p.text:
            continue
        runs = p.runs
        if not runs:
            continue
        giu = runs[0]
        giu.text = p.text.replace("Tên đề tài", ND.TEN_DE_TAI)
        for r in runs[1:]:
            r._r.getparent().remove(r._r)
        giu.font.name = FONT
        rpr = giu._r.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.insert(0, rf)
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rf.set(qn(a), FONT)


# ═══════════════════════════════════════════════════════════════════════════════
#  BẢNG CÓ SẴN TRONG MẪU
# ═══════════════════════════════════════════════════════════════════════════════

def dinh_dang_o(cell, text, bold=False, canh_giua=False, size=Pt(12)):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if canh_giua else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    dong = str(text).split("\n")
    for j, d in enumerate(dong):
        pp = p if j == 0 else cell.add_paragraph()
        if j > 0:
            pp.alignment = p.alignment
            pp.paragraph_format.line_spacing = 1.15
            pp.paragraph_format.space_before = Pt(0)
            pp.paragraph_format.space_after = Pt(2)
        r = pp.add_run(d)
        dat_font(r, size=size, bold=bold)


def dien_bang_co_san(bang, headers, rows, widths_cm, mau_tieu_de="EDF0F5"):
    bang.style = "Table Grid"
    bang.alignment = WD_TABLE_ALIGNMENT.CENTER
    bang.autofit = False
    tblpr = bang._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblpr.append(layout)

    # xóa toàn bộ hàng cũ
    for r in list(bang.rows):
        r._tr.getparent().remove(r._tr)

    hdr = bang.add_row()
    for i, h in enumerate(headers):
        dinh_dang_o(hdr.cells[i], h, bold=True, canh_giua=True)
        to_mau_o(hdr.cells[i], mau_tieu_de)
        hdr.cells[i].width = Cm(widths_cm[i])
    lap_lai_hang_tieu_de(hdr)
    khong_tach_hang(hdr)

    for data in rows:
        row = bang.add_row()
        for i, val in enumerate(data):
            dinh_dang_o(row.cells[i], val, canh_giua=(i == 0))
            row.cells[i].width = Cm(widths_cm[i])
        khong_tach_hang(row)
    for i, w in enumerate(widths_cm):
        bang.columns[i].width = Cm(w)


# ═══════════════════════════════════════════════════════════════════════════════
#  THỰC THI
# ═══════════════════════════════════════════════════════════════════════════════

chuan_hoa_styles()
sua_giao_dien_muc_luc()
chuan_hoa_trang()
don_dep()
sua_trang_bia()
sua_giao_dien_trang_bia()

# Giữ tham chiếu tới hai bảng có sẵn của mẫu TRƯỚC khi chèn thêm bảng mới,
# vì chỉ số trong doc.tables sẽ thay đổi sau mỗi lần chèn.
BANG_VIET_TAT = doc.tables[1]
BANG_PHAN_CONG = doc.tables[2]

# ── Lời cảm ơn ────────────────────────────────────────────────────────────────
chen_sau("LỜI CẢM ƠN", ND.LOI_CAM_ON, chuong=0)

# ── Danh mục ký hiệu, chữ viết tắt ────────────────────────────────────────────
h = tim_heading("DANH MỤC CÁC KÝ HIỆU")
cur = ConTro(h._p)
them_doan(cur, "Bảng dưới đây liệt kê các ký hiệu, chữ viết tắt và thuật ngữ chuyên ngành được sử dụng "
               "xuyên suốt báo cáo, sắp xếp theo trình tự xuất hiện trong nội dung.")
dien_bang_co_san(
    BANG_VIET_TAT,
    ["STT", "Từ viết tắt, ký hiệu", "Viết đầy đủ", "Ý nghĩa"],
    [[str(i + 1), a, b, c] for i, (a, b, c) in enumerate(ND.VIET_TAT)],
    [1.2, 3.0, 4.6, 7.2],
)

# ── Danh mục bảng / hình (trường TOC theo nhãn chú thích) ─────────────────────
h = tim_heading("DANH MỤC CÁC BẢNG")
p = doc.add_paragraph()
h._p.addnext(p._p)
them_field(p, ' TOC \\h \\z \\c "Bang" ',
           cached="Danh mục bảng sẽ hiển thị sau khi cập nhật trường (Ctrl+A rồi F9).")

h = tim_heading("DANH MỤC CÁC HÌNH")
p = doc.add_paragraph()
h._p.addnext(p._p)
them_field(p, ' TOC \\h \\z \\c "Hinh" ',
           cached="Danh mục hình sẽ hiển thị sau khi cập nhật trường (Ctrl+A rồi F9).")

# ── Chương 1 ──────────────────────────────────────────────────────────────────
chen_sau("1.1.", ND.C1_1, chuong=1)
chen_sau("1.2.", ND.C1_2)

# ── Chương 2 ──────────────────────────────────────────────────────────────────
chen_sau("2.1.", ND.C2_1, chuong=2)
chen_sau("2.2.", ND.C2_2)
chen_sau("2.3.", ND.C2_3)

# ── Chương 3 ──────────────────────────────────────────────────────────────────
chen_sau("3.1.", ND.C3_1, chuong=3)
chen_sau("3.2. Xây dựng chức năng Electronic Selling", ND.C3_2_0)
chen_sau("3.2.1.", ND.C3_2_1)
chen_sau("3.2.2.", ND.C3_2_2)
chen_sau("3.2.3.", ND.C3_2_3)
chen_sau("3.3. Xây dựng chức năng Electronic Marketing", ND.C3_3_0)
chen_sau("3.3.1.", ND.C3_3_1)
chen_sau("3.3.2.", ND.C3_3_2)
chen_sau("3.3.3.", ND.C3_3_3)
chen_sau("3.3.4.", ND.C3_3_4)
chen_sau("3.4. Tích hợp", ND.C3_4_0)
chen_sau("3.4.1.", ND.C3_4_1)
chen_sau("3.4.2.", ND.C3_4_2)
chen_sau("3.4.3.", ND.C3_4_3)
chen_sau("3.4.4.", ND.C3_4_4)
chen_sau("3.4.5.", ND.C3_4_5)
chen_sau("3.4.6.", ND.C3_4_6)
chen_sau("3.5. Bảo mật", ND.C3_5_0)
chen_sau("3.5.1.", ND.C3_5_1)
chen_sau("3.5.2.", ND.C3_5_2)
chen_sau("3.5.3.", ND.C3_5_3)
chen_sau("3.5.4.", ND.C3_5_4)
chen_sau("3.6. Triển khai", ND.C3_6_0)
chen_sau("3.6.1.", ND.C3_6_1)
chen_sau("3.6.2.", ND.C3_6_2)
chen_sau("3.6.3.", ND.C3_6_3)
chen_sau("3.6.4.", ND.C3_6_4)

# ── Chương 4 ──────────────────────────────────────────────────────────────────
chen_sau("4.1.", ND.C4_1, chuong=4)
chen_sau("4.2.", ND.C4_2)
chen_sau("4.3.", ND.C4_3)
chen_sau("4.4.", ND.C4_4)

# Bảng phân công (bảng thứ 3 của mẫu) + chú thích ở trên
bang_pc = BANG_PHAN_CONG
p_ct = doc.add_paragraph(style="Caption")
r = p_ct.add_run("Bảng 4.")
dat_font(r, size=Pt(12), italic=True)
them_field(p_ct, " SEQ Bang \\* ARABIC ", cached="2", size=Pt(12), italic=True)
r2 = p_ct.add_run(". Phân công công việc của các thành viên")
dat_font(r2, size=Pt(12), italic=True)
p_ct.paragraph_format.keep_with_next = True
bang_pc._tbl.addprevious(p_ct._p)
dien_bang_co_san(
    bang_pc,
    ["STT", "Thành viên", "Công việc"],
    ND.PHAN_CONG,
    [1.2, 4.8, 10.0],
)

# ── Tài liệu tham khảo ────────────────────────────────────────────────────────
h = tim_heading("TÀI LIỆU THAM KHẢO")
cur = ConTro(h._p)
them_doan(cur, "Danh mục tài liệu tham khảo được trình bày theo chuẩn trích dẫn IEEE.", thut=False)
for i, tl in enumerate(ND.TAI_LIEU, start=1):
    p = p_moi(cur)
    p.paragraph_format.left_indent = Cm(1.1)
    p.paragraph_format.first_line_indent = Cm(-1.1)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"[{i}]\t{tl}")
    dat_font(r)

ep_font_toan_tai_lieu()
bat_cap_nhat_field()
doc.save(OUT)
print("Đã tạo:", OUT)
